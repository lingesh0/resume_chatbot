import os
import spacy
import fitz  # PyMuPDF for PDF text extraction
import gradio as gr
import google.generativeai as genai
from docx import Document
import random
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

# Configure Google Gemini API
genai.configure(api_key="AIzaSyBIp7DNNfyAFxUQKiycBdRadMfvDZ06wmU")  # Replace with your API key

# Load NLP model for ATS Score Calculation
nlp = spacy.load("en_core_web_sm")

# Available resume templates
RESUME_TEMPLATES = {
    "modern": {
        "name": "Modern Professional",
        "description": "Clean and contemporary design with a touch of color",
        "accent_color": RGBColor(79, 70, 229)  # Purple
    },
    "classic": {
        "name": "Classic Executive",
        "description": "Traditional and elegant layout for experienced professionals",
        "accent_color": RGBColor(3, 105, 161)  # Dark Blue
    },
    "creative": {
        "name": "Creative Portfolio",
        "description": "Vibrant and unique design for creative industries",
        "accent_color": RGBColor(236, 72, 153)  # Pink
    },
    "minimal": {
        "name": "Minimalist",
        "description": "Clean, simple layout with focus on content",
        "accent_color": RGBColor(16, 185, 129)  # Green
    },
    "tech": {
        "name": "Tech Professional",
        "description": "Modern design optimized for technical roles",
        "accent_color": RGBColor(245, 158, 11)  # Amber
    }
}

# ✅ Function to Interact with Google Gemini API
def chatbot_interface(user_query, history=[]):
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(user_query)
    history.append((user_query, response.text))
    return history, ""  # Clear input after sending message

# ✅ Enhanced ATS Score Calculation with keyword matching
def calculate_ats_score(resume_text, job_description):
    if not job_description.strip():
        return "⚠️ Please provide a job description to calculate ATS score.", 0.0, {}
    
    resume_doc = nlp(resume_text.lower())
    job_doc = nlp(job_description.lower())
    
    # Extract all meaningful words (excluding stopwords)
    resume_words = {token.lemma_ for token in resume_doc if token.is_alpha and not token.is_stop}
    job_words = {token.lemma_ for token in job_doc if token.is_alpha and not token.is_stop}
    
    if not job_words:
        return "⚠️ Job description is too short to analyze.", 0.0, {}
    
    # Calculate basic match score
    matched_words = resume_words & job_words
    match_score = (len(matched_words) / len(job_words)) * 100
    
    # Identify key skills from job description (more emphasis on nouns and proper nouns)
    key_skills = {token.lemma_: token.text for token in job_doc 
                 if (token.pos_ in ["NOUN", "PROPN"]) and not token.is_stop}
    
    # Calculate skill match rate
    skill_matches = {lemma: orig for lemma, orig in key_skills.items() if lemma in resume_words}
    skill_match_rate = len(skill_matches) / len(key_skills) if key_skills else 0
    
    # Keywords analysis for detailed feedback
    keyword_analysis = {
        "matched_keywords": list(skill_matches.values()),
        "missing_keywords": [key_skills[k] for k in key_skills if k not in resume_words],
        "match_percentage": round(skill_match_rate * 100, 2)
    }
    
    # Final weighted score (basic match + skill match with higher weight)
    final_score = (match_score * 0.4) + (skill_match_rate * 100 * 0.6)
    
    return f"📊 ATS Score: {round(final_score, 2)}%", round(final_score, 2), keyword_analysis

# ✅ AI Suggestions for Resume Improvement with enhanced prompting
def ai_suggestions(resume_text, job_description, keyword_analysis=None):
    if not job_description.strip():
        return "⚠️ Please provide a job description for AI suggestions."
    
    keyword_info = ""
    if keyword_analysis:
        keyword_info = f"""
        Keyword Analysis:
        - Matched Keywords: {', '.join(keyword_analysis.get('matched_keywords', []))}
        - Missing Keywords: {', '.join(keyword_analysis.get('missing_keywords', []))}
        - Keyword Match Rate: {keyword_analysis.get('match_percentage', 0)}%
        """
    
    prompt = f"""
    You are an expert resume consultant specializing in ATS optimization.
    
    Analyze this resume and provide detailed suggestions for improvement based on the given job description.
    
    Resume Content: {resume_text}
    
    Job Description: {job_description}
    
    {keyword_info}
    
    Provide a comprehensive, actionable analysis with the following sections:
    
    1. KEYWORD OPTIMIZATION:
       - Identify critical keywords from the job description that should be added to the resume
       - Suggest specific placements or phrasings to naturally incorporate these keywords
    
    2. SKILLS GAP ANALYSIS:
       - Identify skills mentioned in the job that aren't highlighted in the resume
       - Suggest ways to address or frame existing experience to cover these gaps
    
    3. CONTENT STRUCTURE RECOMMENDATIONS:
       - Analyze if the resume layout and section organization is optimal
       - Suggest improvements to section ordering, bullet point structure, etc.
       
    4. ACHIEVEMENT EMPHASIS:
       - Suggest how to transform job descriptions into achievement statements
       - Recommend how to quantify accomplishments for greater impact
    
    5. TOP 3 PRIORITY CHANGES:
       - List the three most critical changes needed for maximum ATS score improvement
       
    Format your response in markdown with clear headings and bullet points.
    """
    
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(prompt)
    return response.text

# ✅ Function to extract skills from job description
def extract_key_skills(job_description):
    if not job_description.strip():
        return []
    
    prompt = f"""
    Extract and list all hard skills, technical skills, soft skills, and qualifications mentioned 
    in this job description. Format as a JSON array of strings with just the skill names.
    
    Job Description:
    {job_description}
    
    Response format example:
    ["Python", "Project Management", "Communication", "Bachelor's degree"]
    """
    
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(prompt)
    
    try:
        # Try to parse as JSON
        skills = json.loads(response.text)
        if isinstance(skills, list):
            return skills
        return []
    except:
        # Fallback: try to extract skills from text response
        skills_text = response.text.strip("` \n")
        if skills_text.startswith("[") and skills_text.endswith("]"):
            try:
                return json.loads(skills_text)
            except:
                pass
        
        # Last resort: split by comma or newline
        if "," in skills_text:
            return [s.strip() for s in skills_text.split(",")]
        elif "\n" in skills_text:
            return [s.strip() for s in skills_text.split("\n") if s.strip()]
        
        return []

# ✅ Enhanced function to generate tailored resume sections from job description
def generate_tailored_sections(resume_data, job_description):
    if not job_description.strip():
        return None, None
    
    # Create a compact version of the resume data for the API call
    resume_summary = {
        "full_name": resume_data.get("full_name", ""),
        "education": resume_data.get("education", {}),
        "skills": resume_data.get("skills", []),
        "experience": resume_data.get("experience", []),
        "projects": resume_data.get("projects", [])
    }
    
    prompt = f"""
    As an AI resume expert, create a tailored professional summary and personal objective 
    for a resume based on the candidate's information and the job description.
    
    CANDIDATE INFORMATION:
    {json.dumps(resume_summary, indent=2)}
    
    JOB DESCRIPTION:
    {job_description}
    
    Respond with ONLY a JSON object with two properties:
    1. "professional_summary": A powerful 3-4 sentence summary highlighting the candidate's most relevant experience, skills, and qualifications for this specific role
    2. "career_objective": A concise 1-2 sentence statement about the candidate's career goals aligned with this specific position
    
    Example format:
    {{
      "professional_summary": "...",
      "career_objective": "..."
    }}
    """
    
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(prompt)
    
    try:
        result = json.loads(response.text)
        return result.get("professional_summary"), result.get("career_objective")
    except:
        # Fallback method in case JSON parsing fails
        text = response.text
        summary = ""
        objective = ""
        
        if "professional_summary" in text.lower():
            summary_parts = text.lower().split("professional_summary")
            if len(summary_parts) > 1:
                summary_text = summary_parts[1].split("career_objective")[0]
                summary = summary_text.strip().strip('":,{}\n')
        
        if "career_objective" in text.lower():
            objective_parts = text.lower().split("career_objective")
            if len(objective_parts) > 1:
                objective_text = objective_parts[1].strip()
                objective = objective_text.strip('":,{}\n')
        
        return summary, objective

# ✅ Enhanced function to generate resume with improved formatting and template options
def generate_resume(resume_data, job_description, template="modern"):
    if not resume_data.get("full_name"):
        return None, "⚠️ Please provide at least a name to generate a resume.", "", {}
    
    # Get template configuration
    template_config = RESUME_TEMPLATES.get(template, RESUME_TEMPLATES["modern"])
    accent_color = template_config["accent_color"]
    
    # Create a new document
    doc = Document()
    
    # Set up document style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Generate tailored sections based on job description
    prof_summary, career_objective = generate_tailored_sections(resume_data, job_description)
    
    # Apply template-specific styling
    if template == "modern":
        # Add name as header
        name_heading = doc.add_heading('', level=0)
        name_run = name_heading.add_run(resume_data.get("full_name", "").upper())
        name_run.font.size = Pt(24)
        name_run.font.color.rgb = accent_color
        name_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add contact information
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_info = []
        if resume_data.get("email"):
            contact_info.append(resume_data.get("email"))
        if resume_data.get("phone"):
            contact_info.append(resume_data.get("phone"))
        if resume_data.get("location"):
            contact_info.append(resume_data.get("location"))
        if resume_data.get("linkedin"):
            contact_info.append(resume_data.get("linkedin"))
        
        contact_run = contact_para.add_run(' | '.join(contact_info))
        contact_run.font.size = Pt(11)
        contact_run.italic = True
        
        # Add horizontal line
        doc.add_paragraph('_' * 100)
        
    elif template == "classic":
        # Header with name and contact side by side
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        
        # Name cell
        name_cell = table.cell(0, 0)
        name_para = name_cell.paragraphs[0]
        name_run = name_para.add_run(resume_data.get("full_name", ""))
        name_run.bold = True
        name_run.font.size = Pt(18)
        
        # Contact cell
        contact_cell = table.cell(0, 1)
        contact_para = contact_cell.paragraphs[0]
        contact_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if resume_data.get("phone"):
            contact_para.add_run(f"Phone: {resume_data.get('phone')}\n")
        if resume_data.get("email"):
            contact_para.add_run(f"Email: {resume_data.get('email')}\n")
        if resume_data.get("location"):
            contact_para.add_run(f"Location: {resume_data.get('location')}\n")
        if resume_data.get("linkedin"):
            contact_para.add_run(f"LinkedIn: {resume_data.get('linkedin')}")
            
    elif template == "creative":
        # Add a colored rectangle at the top
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture('rect.png', width=Inches(6))  # This is a placeholder - we can't actually add an image here
        
        # Creative spacing
        doc.add_paragraph()
        
        # Add name as large header
        name_heading = doc.add_heading('', level=0)
        name_run = name_heading.add_run(resume_data.get("full_name", ""))
        name_run.font.size = Pt(28)
        name_run.font.color.rgb = accent_color
        
        # Add title if available
        if resume_data.get("desired_position"):
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(resume_data.get("desired_position"))
            title_run.font.size = Pt(16)
            title_run.italic = True
        
        # Add contact in creative format
        contact_para = doc.add_paragraph()
        contact_para.add_run("📧 ").bold = True
        contact_para.add_run(resume_data.get("email", ""))
        contact_para.add_run(" • 📱 ").bold = True
        contact_para.add_run(resume_data.get("phone", ""))
        
    elif template in ["minimal", "tech"]:
        # Minimal header
        name_heading = doc.add_heading('', level=0)
        name_run = name_heading.add_run(resume_data.get("full_name", ""))
        name_run.font.size = Pt(20)
        
        # Simple contact line
        contact_para = doc.add_paragraph()
        contacts = []
        if resume_data.get("email"):
            contacts.append(resume_data.get("email"))
        if resume_data.get("phone"):
            contacts.append(resume_data.get("phone"))
        if resume_data.get("location"):
            contacts.append(resume_data.get("location"))
        
        contact_para.add_run(" • ".join(contacts))
        
        # Tech template adds a skill progress bar visualization (simulated)
        if template == "tech":
            doc.add_paragraph()
            tech_para = doc.add_paragraph()
            tech_para.add_run("Tech stack: ").bold = True
            tech_para.add_run(" | ".join(resume_data.get("skills", [])[:5]))
    
    # Add Professional Summary section if available
    if prof_summary:
        doc.add_paragraph()
        summary_heading = doc.add_heading('PROFESSIONAL SUMMARY', level=1)
        summary_heading.style.font.color.rgb = accent_color
        summary_para = doc.add_paragraph()
        summary_para.add_run(prof_summary)
    
    # Add Career Objective if available
    if career_objective:
        doc.add_paragraph()
        objective_heading = doc.add_heading('CAREER OBJECTIVE', level=1)
        objective_heading.style.font.color.rgb = accent_color
        objective_para = doc.add_paragraph()
        objective_para.add_run(career_objective)
    
    # Add Education section
    if resume_data.get("education"):
        doc.add_paragraph()
        education_heading = doc.add_heading('EDUCATION', level=1)
        education_heading.style.font.color.rgb = accent_color
        
        for edu in resume_data.get("education", []):
            edu_para = doc.add_paragraph()
            if edu.get("degree"):
                edu_run = edu_para.add_run(f"{edu.get('degree')}")
                edu_run.bold = True
            
            if edu.get("institution"):
                if edu.get("degree"):
                    edu_para.add_run(f", {edu.get('institution')}")
                else:
                    edu_run = edu_para.add_run(f"{edu.get('institution')}")
                    edu_run.bold = True
            
            # Add date range on the same line but right-aligned
            if edu.get("year_start") and edu.get("year_end"):
                date_text = f" ({edu.get('year_start')} - {edu.get('year_end')})"
                edu_para.add_run(date_text)
            
            # Add GPA/percentage if available
            if edu.get("gpa") or edu.get("percentage"):
                gpa_para = doc.add_paragraph()
                gpa_para.style = 'List Bullet'
                if edu.get("gpa"):
                    gpa_para.add_run(f"GPA: {edu.get('gpa')}")
                if edu.get("percentage"):
                    gpa_para.add_run(f"Percentage: {edu.get('percentage')}%")
            
            # Add relevant coursework if available
            if edu.get("coursework"):
                course_para = doc.add_paragraph()
                course_para.style = 'List Bullet'
                course_para.add_run(f"Relevant Coursework: {edu.get('coursework')}")
    
    # Add Skills section
    if resume_data.get("skills"):
        doc.add_paragraph()
        skills_heading = doc.add_heading('SKILLS', level=1)
        skills_heading.style.font.color.rgb = accent_color
        
        # Organize skills by categories if available
        skills_by_category = {}
        for skill in resume_data.get("skills", []):
            category = "Technical Skills"  # Default category
            if isinstance(skill, dict) and skill.get("category") and skill.get("name"):
                category = skill.get("category")
                skill_name = skill.get("name")
            else:
                skill_name = skill if isinstance(skill, str) else str(skill)
            
            if category not in skills_by_category:
                skills_by_category[category] = []
            skills_by_category[category].append(skill_name)
        
        # Add each skill category
        for category, skills_list in skills_by_category.items():
            if skills_list:
                cat_para = doc.add_paragraph()
                cat_para.add_run(f"{category}: ").bold = True
                cat_para.add_run(", ".join(skills_list))
    
    # Add Experience section
    if resume_data.get("experience"):
        doc.add_paragraph()
        exp_heading = doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
        exp_heading.style.font.color.rgb = accent_color
        
        for exp in resume_data.get("experience", []):
            exp_para = doc.add_paragraph()
            
            # Add position
            if exp.get("position"):
                position_run = exp_para.add_run(exp.get("position"))
                position_run.bold = True
            
            # Add company
            if exp.get("company"):
                if exp.get("position"):
                    exp_para.add_run(f", {exp.get('company')}")
                else:
                    company_run = exp_para.add_run(exp.get("company"))
                    company_run.bold = True
            
            # Add date range
            if exp.get("start_date") and exp.get("end_date"):
                exp_para.add_run(f" ({exp.get('start_date')} - {exp.get('end_date')})")
            elif exp.get("start_date"):
                exp_para.add_run(f" ({exp.get('start_date')} - Present)")
            
            # Add location if available
            if exp.get("location"):
                loc_para = doc.add_paragraph()
                loc_para.add_run(f"Location: {exp.get('location')}")
            
            # Add responsibilities and achievements
            if exp.get("responsibilities"):
                for resp in exp.get("responsibilities"):
                    resp_para = doc.add_paragraph()
                    resp_para.style = 'List Bullet'
                    resp_para.add_run(resp)
    
    # Add Projects section
    if resume_data.get("projects"):
        doc.add_paragraph()
        proj_heading = doc.add_heading('PROJECTS', level=1)
        proj_heading.style.font.color.rgb = accent_color
        
        for project in resume_data.get("projects", []):
            proj_para = doc.add_paragraph()
            
            # Add project name
            if project.get("name"):
                proj_run = proj_para.add_run(project.get("name"))
                proj_run.bold = True
            
            # Add date or duration
            if project.get("date"):
                proj_para.add_run(f" ({project.get('date')})")
            
            # Add project description
            if project.get("description"):
                desc_para = doc.add_paragraph()
                desc_para.style = 'List Bullet'
                desc_para.add_run(project.get("description"))
            
            # Add technologies used
            if project.get("technologies"):
                tech_para = doc.add_paragraph()
                tech_para.style = 'List Bullet'
                tech_para.add_run(f"Technologies: {project.get('technologies')}")
            
            # Add project URL if available
            if project.get("url"):
                url_para = doc.add_paragraph()
                url_para.style = 'List Bullet'
                url_para.add_run(f"URL: {project.get('url')}")
    
    # Add Certifications section if available
    if resume_data.get("certifications"):
        doc.add_paragraph()
        cert_heading = doc.add_heading('CERTIFICATIONS', level=1)
        cert_heading.style.font.color.rgb = accent_color
        
        for cert in resume_data.get("certifications", []):
            cert_para = doc.add_paragraph()
            cert_para.style = 'List Bullet'
            
            if isinstance(cert, dict):
                cert_text = cert.get("name", "")
                if cert.get("issuer"):
                    cert_text += f" - {cert.get('issuer')}"
                if cert.get("date"):
                    cert_text += f" ({cert.get('date')})"
                cert_para.add_run(cert_text)
            else:
                cert_para.add_run(str(cert))
    
    # Add Languages section if available
    if resume_data.get("languages"):
        doc.add_paragraph()
        lang_heading = doc.add_heading('LANGUAGES', level=1)
        lang_heading.style.font.color.rgb = accent_color
        
        lang_para = doc.add_paragraph()
        lang_list = []
        
        for lang in resume_data.get("languages", []):
            if isinstance(lang, dict) and lang.get("name") and lang.get("proficiency"):
                lang_list.append(f"{lang.get('name')} ({lang.get('proficiency')})")
            elif isinstance(lang, str):
                lang_list.append(lang)
        
        lang_para.add_run(", ".join(lang_list))
    
    # Add Interests/Hobbies section if available
    if resume_data.get("interests"):
        doc.add_paragraph()
        interests_heading = doc.add_heading('INTERESTS', level=1)
        interests_heading.style.font.color.rgb = accent_color
        
        interests_para = doc.add_paragraph()
        
        if isinstance(resume_data.get("interests"), list):
            interests_para.add_run(", ".join(resume_data.get("interests")))
        else:
            interests_para.add_run(str(resume_data.get("interests")))
    
    # Save the document
    filename = f"resume_{resume_data.get('full_name', 'user').replace(' ', '_')}.docx"
    doc.save(filename)
    
    # Calculate ATS score
    resume_text = " ".join([
        resume_data.get("full_name", ""),
        " ".join(str(s) for s in resume_data.get("skills", [])),
        " ".join(str(e.get("responsibilities", "")) for e in resume_data.get("experience", [])),
        " ".join(str(p.get("description", "")) for p in resume_data.get("projects", []))
    ])
    
    ats_text, score, keyword_analysis = calculate_ats_score(resume_text, job_description)
    ai_suggestion_text = ai_suggestions(resume_text, job_description, keyword_analysis)
    
    return filename, ats_text, ai_suggestion_text, score

# ✅ Function to Analyze Uploaded Resume
def analyze_uploaded_resume(uploaded_resume, job_description):
    if not uploaded_resume or not job_description.strip():
        return "⚠️ Please upload a resume and provide a job description.", None, 0.0
    
    resume_text = ""
    file_ext = uploaded_resume.name.split(".")[-1].lower()
    
    try:
        if file_ext == "docx":
            doc = Document(uploaded_resume.name)
            resume_text = "\n".join([para.text for para in doc.paragraphs])
        elif file_ext == "pdf":
            with fitz.open(uploaded_resume.name) as pdf_doc:
                resume_text = "\n".join([page.get_text("text") for page in pdf_doc])
        else:
            return "❌ Unsupported file format! Please upload a DOCX or PDF file.", None, 0.0
        
        ats_text, score, keyword_analysis = calculate_ats_score(resume_text, job_description)
        ai_suggestion_text = ai_suggestions(resume_text, job_description, keyword_analysis)
        return ats_text, ai_suggestion_text, score
    except Exception as e:
        return f"❌ Error processing resume: {str(e)}", None, 0.0

# Function to process form data into structured format
def process_resume_form(full_name, user_email, contact_number, location, linkedin, education_degree, 
                        education_institution, education_start, education_end, education_gpa, 
                        education_percentage, education_coursework, user_skills, 
                        experience_company, experience_position, experience_start_date, 
                        experience_end_date, experience_location, experience_responsibilities,
                        project_name, project_date, project_description, project_technologies, 
                        project_url, certifications, languages, interests,
                        desired_position, job_description):
    
    # Create structured education data
    education = []
    for i in range(len(education_institution)):
        if education_institution[i]:
            edu = {
                "institution": education_institution[i],
                "degree": education_degree[i] if i < len(education_degree) else "",
                "year_start": education_start[i] if i < len(education_start) else "",
                "year_end": education_end[i] if i < len(education_end) else "",
                "gpa": education_gpa[i] if i < len(education_gpa) else "",
                "percentage": education_percentage[i] if i < len(education_percentage) else "",
                "coursework": education_coursework[i] if i < len(education_coursework) else ""
            }
            education.append(edu)
    
    # Parse skills
    skills = [skill.strip() for skill in user_skills.split(',') if skill.strip()]
    
    # Add extracted skills from job description if available
    if job_description:
        extracted_skills = extract_key_skills(job_description)
        for skill in extracted_skills:
            if skill not in skills:
                # Add with special category to highlight
                skills.append({"name": skill, "category": "Job-Relevant Skills"})
    
    # Create structured experience data
    experience = []
    for i in range(len(experience_company)):
        if experience_company[i] and experience_position[i]:
            exp = {
                "company": experience_company[i],
                "position": experience_position[i],
                "start_date": experience_start_date[i] if i < len(experience_start_date) else "",
                "end_date": experience_end_date[i] if i < len(experience_end_date) else "",
                "location": experience_location[i] if i < len(experience_location) else "",
                "responsibilities": experience_responsibilities[i].split('\n') if i < len(experience_responsibilities) and experience_responsibilities[i] else []
            }
            experience.append(exp)
    
    # Create structured project data
    projects = []
    for i in range(len(project_name)):
        if project_name[i]:
            proj = {
                "name": project_name[i],
                "date": project_date[i] if i < len(project_date) else "",
                "description": project_description[i] if i < len(project_description) else "",
                "technologies": project_technologies[i] if i < len(project_technologies) else "",
                "url": project_url[i] if i < len(project_url) else ""
            }
            projects.append(proj)
    
    # Parse certifications
    cert_list = [cert.strip() for cert in certifications.split(',') if cert.strip()]
    
    # Parse languages
    lang_list = [lang.strip() for lang in languages.split(',') if lang.strip()]
    
    # Parse interests
    interest_list = [interest.strip() for interest in interests.split(',') if interest.strip()]
    
    # Create the complete resume data structure
    resume_data = {
        "full_name": full_name,
        "email": user_email,
        "phone": contact_number,
        "location": location,
        "linkedin": linkedin,
        "education": education,
        "skills": skills,
        "experience": experience,
        "projects": projects,
        "certifications": cert_list,
        "languages": lang_list,
        "interests": interest_list,
        "desired_position": desired_position
    }
    
    return resume_data

# Enhanced CSS with animations and better styling
custom_css = """
:root {
    --primary-color: #4F46E5;
    --secondary-color: #A78BFA;
    --success-color: #10B981;
    --warning-color: #F59E0B;
    --danger-color: #EF4444;
    --bg-color: #0F172A;
    --card-bg: #1E293B;
    --card-bg-light: #334155;
    --text-color: #F1F5F9;
    --text-muted: #CBD5E1;
    --border-color: #475569;
    --border-radius: 16px;
    --box-shadow: 0 25px 50px -12px rgba(0, 0, 0, 0.25);
    --transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
}

* {
    transition: var(--transition);
}

body {
    background: linear-gradient(135deg, #0F172A 0%, #1E293B 50%, #334155 100%) !important;
    color: var(--text-color) !important;
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif !important;
    min-height: 100vh !important;
}

.gradio-container {
    width: 95% !important;
    max-width: 1400px !important;
    margin: auto !important;
    padding: 20px !important;
}

.container {
    background: rgba(30, 41, 59, 0.8) !important;
    backdrop-filter: blur(10px) !important;
    border: 1px solid rgba(71, 85, 105, 0.3) !important;
    border-radius: var(--border-radius) !important;
    padding: 32px !important;
    margin-bottom: 24px !important;
    box-shadow: var(--box-shadow) !important;
    position: relative !important;
    overflow: hidden !important;
}

.container::before {
    content: '' !important;
    position: absolute !important;
    top: 0 !important;
    left: 0 !important;
    right: 0 !important;
    height: 2px !important;
    background: linear-gradient(90deg, var(--primary-color), var(--secondary-color)) !important;
}

.container:hover {
    transform: translateY(-8px) !important;
    box-shadow: 0 32px 64px -12px rgba(79, 70, 229, 0.2) !important;
    border-color: rgba(79, 70, 229, 0.5) !important;
}

.tabs {
    border-radius: var(--border-radius) !important;
    overflow: hidden !important;
    background: var(--card-bg) !important;
    box-shadow: var(--box-shadow) !important;
}

.tab-nav {
    background: var(--card-bg) !important;
    border-bottom: 1px solid var(--border-color) !important;
    padding: 0 !important;
}

.tab-nav button {
    font-weight: 600 !important;
    font-size: 1rem !important;
    padding: 16px 24px !important;
    border: none !important;
    background: transparent !important;
    color: var(--text-muted) !important;
    position: relative !important;
    cursor: pointer !important;
}

.tab-nav button::before {
    content: '' !important;
    position: absolute !important;
    bottom: 0 !important;
    left: 50% !important;
    width: 0 !important;
    height: 3px !important;
    background: linear-gradient(90deg, var(--primary-color), var(--secondary-color)) !important;
    transform: translateX(-50%) !important;
    transition: var(--transition) !important;
}

.tab-nav button:hover::before {
    width: 50% !important;
}

.tab-nav button.selected {
    background: linear-gradient(135deg, var(--primary-color) 0%, var(--secondary-color) 100%) !important;
    color: white !important;
}

.tab-nav button.selected::before {
    width: 100% !important;
}

input, textarea, select {
    background: var(--card-bg-light) !important;
    border: 2px solid var(--border-color) !important;
    color: var(--text-color) !important;
    border-radius: 12px !important;
    padding: 14px 18px !important;
    font-size: 1rem !important;
    font-weight: 500 !important;
}

input:focus, textarea:focus, select:focus {
    border-color: var(--primary-color) !important;
    box-shadow: 0 0 0 4px rgba(79, 70, 229, 0.1) !important;
    outline: none !important;
    transform: translateY(-2px) !important;
}

input::placeholder, textarea::placeholder {
    color: var(--text-muted) !important;
    opacity: 0.7 !important;
}

button {
    font-weight: 600 !important;
    border-radius: 12px !important;
    padding: 14px 28px !important;
    border: none !important;
    cursor: pointer !important;
    font-size: 1rem !important;
    position: relative !important;
    overflow: hidden !important;
}

button.primary {
    background: linear-gradient(135deg, var(--primary-color) 0%, var(--secondary-color) 100%) !important;
    color: white !important;
    box-shadow: 0 8px 25px -8px var(--primary-color) !important;
}

button.primary::before {
    content: '' !important;
    position: absolute !important;
    top: 0 !important;
    left: -100% !important;
    width: 100% !important;
    height: 100% !important;
    background: linear-gradient(90deg, transparent, rgba(255,255,255,0.2), transparent) !important;
    transition: left 0.5s !important;
}

button.primary:hover::before {
    left: 100% !important;
}

button.primary:hover {
    transform: translateY(-3px) !important;
    box-shadow: 0 12px 35px -8px var(--primary-color) !important;
}

button.secondary {
    background: var(--card-bg-light) !important;
    color: var(--text-color) !important;
    border: 2px solid var(--border-color) !important;
}

button.secondary:hover {
    background: var(--primary-color) !important;
    border-color: var(--primary-color) !important;
    transform: translateY(-2px) !important;
}

.score-display {
    font-size: 3rem !important;
    font-weight: 800 !important;
    text-align: center !important;
    padding: 32px !important;
    border-radius: var(--border-radius) !important;
    margin: 20px 0 !important;
    position: relative !important;
    overflow: hidden !important;
}

.score-display::before {
    content: '' !important;
    position: absolute !important;
    top: -50% !important;
    left: -50% !important;
    width: 200% !important;
    height: 200% !important;
    background: conic-gradient(from 0deg, transparent, currentColor, transparent) !important;
    animation: rotate 3s linear infinite !important;
    opacity: 0.1 !important;
}

@keyframes rotate {
    to { transform: rotate(360deg); }
}

.score-low {
    background: rgba(239, 68, 68, 0.1) !important;
    color: var(--danger-color) !important;
    border: 2px solid rgba(239, 68, 68, 0.3) !important;
}

.score-medium {
    background: rgba(245, 158, 11, 0.1) !important;
    color: var(--warning-color) !important;
    border: 2px solid rgba(245, 158, 11, 0.3) !important;
}

.score-high {
    background: rgba(16, 185, 129, 0.1) !important;
    color: var(--success-color) !important;
    border: 2px solid rgba(16, 185, 129, 0.3) !important;
}

.header {
    background: linear-gradient(135deg, #4F46E5 0%, #A78BFA 25%, #EC4899 50%, #8B5CF6 75%, #4F46E5 100%) !important;
    background-size: 400% 400% !important;
    animation: gradient-shift 8s ease infinite !important;
    padding: 48px 32px !important;
    border-radius: var(--border-radius) !important;
    margin-bottom: 32px !important;
    text-align: center !important;
    box-shadow: var(--box-shadow) !important;
    position: relative !important;
    overflow: hidden !important;
}

.header::before {
    content: '' !important;
    position: absolute !important;
    top: 0 !important;
    left: 0 !important;
    right: 0 !important;
    bottom: 0 !important;
    background: linear-gradient(45deg, transparent 30%, rgba(255,255,255,0.1) 50%, transparent 70%) !important;
    animation: shimmer 3s ease-in-out infinite !important;
}

@keyframes gradient-shift {
    0%, 100% { background-position: 0% 50%; }
    50% { background-position: 100% 50%; }
}

@keyframes shimmer {
    0% { transform: translateX(-100%); }
    100% { transform: translateX(100%); }
}

.header h1 {
    font-size: 3.5rem !important;
    font-weight: 900 !important;
    margin-bottom: 16px !important;
    color: white !important;
    text-shadow: 0 4px 8px rgba(0, 0, 0, 0.3) !important;
    position: relative !important;
    z-index: 1 !important;
}

.header p {
    font-size: 1.25rem !important;
    max-width: 700px !important;
    margin: 0 auto !important;
    color: rgba(255, 255, 255, 0.9) !important;
    font-weight: 500 !important;
    position: relative !important;
    z-index: 1 !important;
}

.form-section {
    background: var(--card-bg) !important;
    border-radius: var(--border-radius) !important;
    padding: 24px !important;
    margin-bottom: 24px !important;
    border: 1px solid var(--border-color) !important;
    position: relative !important;
}

.form-section h3 {
    color: var(--primary-color) !important;
    margin-top: 0 !important;
    margin-bottom: 20px !important;
    font-size: 1.5rem !important;
    font-weight: 700 !important;
    display: flex !important;
    align-items: center !important;
    gap: 12px !important;
}

.form-section h3::before {
    content: '' !important;
    width: 4px !important;
    height: 24px !important;
    background: linear-gradient(135deg, var(--primary-color), var(--secondary-color)) !important;
    border-radius: 2px !important;
}

.template-card {
    background: var(--card-bg-light) !important;
    border: 2px solid var(--border-color) !important;
    border-radius: 12px !important;
    padding: 20px !important;
    cursor: pointer !important;
    transition: var(--transition) !important;
    text-align: center !important;
}

.template-card:hover {
    border-color: var(--primary-color) !important;
    transform: translateY(-4px) !important;
    box-shadow: 0 12px 25px -8px rgba(79, 70, 229, 0.3) !important;
}

.template-card.selected {
    border-color: var(--primary-color) !important;
    background: rgba(79, 70, 229, 0.1) !important;
}

.template-card h4 {
    color: var(--text-color) !important;
    margin-bottom: 8px !important;
    font-weight: 600 !important;
}

.template-card p {
    color: var(--text-muted) !important;
    font-size: 0.9rem !important;
    margin: 0 !important;
}

.add-btn {
    background: var(--success-color) !important;
    color: white !important;
    border: none !important;
    border-radius: 8px !important;
    padding: 8px 16px !important;
    font-size: 0.9rem !important;
    font-weight: 600 !important;
    cursor: pointer !important;
    margin-top: 16px !important;
}

.add-btn:hover {
    background: #059669 !important;
    transform: translateY(-2px) !important;
}

.remove-btn {
    background: var(--danger-color) !important;
    color: white !important;
    border: none !important;
    border-radius: 8px !important;
    padding: 8px 16px !important;
    font-size: 0.9rem !important;
    font-weight: 600 !important;
    cursor: pointer !important;
    margin-left: 8px !important;
}

.remove-btn:hover {
    background: #dc2626 !important;
    transform: translateY(-2px) !important;
}

.chatbot-container {
    height: 600px !important;
    border-radius: var(--border-radius) !important;
    background: var(--card-bg-light) !important;
    border: 1px solid var(--border-color) !important;
    overflow: hidden !important;
}

.suggestion-box {
    background: rgba(79, 70, 229, 0.05) !important;
    border-left: 4px solid var(--primary-color) !important;
    border-radius: 0 12px 12px 0 !important;
    padding: 20px !important;
    margin: 20px 0 !important;
}

.suggestion-box h4 {
    color: var(--primary-color) !important;
    margin-top: 0 !important;
    margin-bottom: 12px !important;
    font-size: 1.2rem !important;
    font-weight: 700 !important;
}

.info-card {
    background: var(--card-bg) !important;
    border-radius: var(--border-radius) !important;
    padding: 24px !important;
    box-shadow: 0 8px 25px -8px rgba(0, 0, 0, 0.3) !important;
    border: 1px solid var(--border-color) !important;
    transition: var(--transition) !important;
}

.info-card:hover {
    transform: translateY(-6px) !important;
    box-shadow: 0 16px 35px -8px rgba(79, 70, 229, 0.2) !important;
}

.info-card h3 {
    color: var(--primary-color) !important;
    margin-top: 0 !important;
    font-size: 1.3rem !important;
    margin-bottom: 12px !important;
    font-weight: 700 !important;
}

.progress-bar {
    width: 100% !important;
    height: 8px !important;
    background: var(--card-bg-light) !important;
    border-radius: 4px !important;
    overflow: hidden !important;
    margin: 12px 0 !important;
}

.progress-fill {
    height: 100% !important;
    background: linear-gradient(90deg, var(--primary-color), var(--secondary-color)) !important;
    transition: width 1s ease !important;
}

.file-upload {
    position: relative !important;
    display: block !important;
    width: 100% !important;
    margin: 16px 0 !important;
}

.file-upload input[type="file"] {
    opacity: 0 !important;
    position: absolute !important;
    width: 100% !important;
    height: 100% !important;
    cursor: pointer !important;
}

.file-upload-label {
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    padding: 32px 16px !important;
    background: var(--card-bg-light) !important;
    border: 3px dashed var(--border-color) !important;
    border-radius: var(--border-radius) !important;
    color: var(--text-color) !important;
    font-weight: 600 !important;
    font-size: 1.1rem !important;
    transition: var(--transition) !important;
    cursor: pointer !important;
    min-height: 120px !important;
    flex-direction: column !important;
    gap: 12px !important;
}

.file-upload-label:hover {
    border-color: var(--primary-color) !important;
    background: rgba(79, 70, 229, 0.05) !important;
    transform: translateY(-2px) !important;
}

.file-upload-icon {
    font-size: 2rem !important;
    opacity: 0.7 !important;
}

.stats-card {
    background: linear-gradient(135deg, var(--primary-color), var(--secondary-color)) !important;
    color: white !important;
    padding: 24px !important;
    border-radius: var(--border-radius) !important;
    text-align: center !important;
    box-shadow: 0 8px 25px -8px var(--primary-color) !important;
}

.stats-number {
    font-size: 2.5rem !important;
    font-weight: 900 !important;
    margin-bottom: 8px !important;
}

.stats-label {
    font-size: 1rem !important;
    opacity: 0.9 !important;
    font-weight: 500 !important;
}

/* Custom scrollbar */
::-webkit-scrollbar {
    width: 12px !important;
    height: 12px !important;
}

::-webkit-scrollbar-track {
    background: var(--card-bg) !important;
    border-radius: 6px !important;
}

::-webkit-scrollbar-thumb {
    background: linear-gradient(135deg, var(--primary-color), var(--secondary-color)) !important;
    border-radius: 6px !important;
    border: 2px solid var(--card-bg) !important;
}

::-webkit-scrollbar-thumb:hover {
    background: linear-gradient(135deg, var(--secondary-color), var(--primary-color)) !important;
}

.tooltip {
    position: relative !important;
    cursor: help !important;
}

.tooltip::after {
    content: attr(data-tooltip) !important;
    position: absolute !important;
    bottom: 125% !important;
    left: 50% !important;
    transform: translateX(-50%) !important;
    background: var(--card-bg) !important;
    color: var(--text-color) !important;
    padding: 8px 12px !important;
    border-radius: 8px !important;
    font-size: 0.9rem !important;
    white-space: nowrap !important;
    opacity: 0 !important;
    pointer-events: none !important;
    transition: opacity 0.3s !important;
    border: 1px solid var(--border-color) !important;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.3) !important;
    z-index: 1000 !important;
}

.tooltip:hover::after {
    opacity: 1 !important;
}

/* Responsive design */
@media (max-width: 768px) {
    .gradio-container {
        width: 98% !important;
        padding: 10px !important;
    }
    
    .container {
        padding: 20px !important;
    }
    
    .header h1 {
        font-size: 2.5rem !important;
    }
    
    .header p {
        font-size: 1.1rem !important;
    }
    
    .score-display {
        font-size: 2rem !important;
        padding: 20px !important;
    }
}

/* Loading animation */
@keyframes pulse {
    0%, 100% { opacity: 1; }
    50% { opacity: 0.5; }
}

.loading {
    animation: pulse 2s infinite !important;
}

/* Success animation */
@keyframes checkmark {
    0% { transform: rotate(45deg) scale(0); }
    100% { transform: rotate(45deg) scale(1); }
}

.success-check {
    animation: checkmark 0.6s ease-in-out !important;
}
"""

# Function to update score display class based on score value
def update_score_class(score):
    if score < 40:
        return "score-low"
    elif score < 70:
        return "score-medium"
    else:
        return "score-high"

# ✅ Enhanced Gradio UI with comprehensive form and styling
with gr.Blocks(css=custom_css, title="AI Resume Builder & ATS Optimizer") as iface:
    # Header section
    gr.HTML("""
    <div class="header">
        <h1>🚀 AI Resume Builder & ATS Optimizer</h1>
        <p>Create professional, ATS-optimized resumes with AI-powered suggestions and multiple templates. 
        Get detailed feedback and land your dream job faster.</p>
    </div>
    """)
    
    with gr.Tabs(elem_classes=["tabs"]) as tabs:
        # Tab 1: Enhanced Resume Generator
        with gr.Tab("📄 Create Professional Resume", elem_id="create-tab"):
            with gr.Row():
                with gr.Column(scale=2):
                    # Personal Information Section
                    with gr.Group(elem_classes=["form-section"]):
                        gr.HTML("<h3>👤 Personal Information</h3>")
                        with gr.Row():
                            full_name = gr.Textbox(
                                label="Full Name *", 
                                placeholder="John Doe",
                                elem_classes=["tooltip"],
                                elem_id="name-input"
                            )
                            user_email = gr.Textbox(
                                label="Email Address *", 
                                placeholder="john.doe@example.com"
                            )
                        
                        with gr.Row():
                            contact_number = gr.Textbox(
                                label="Phone Number *", 
                                placeholder="+1 (555) 123-4567"
                            )
                            location = gr.Textbox(
                                label="Location", 
                                placeholder="City, State/Country"
                            )
                        
                        linkedin = gr.Textbox(
                            label="LinkedIn Profile", 
                            placeholder="https://linkedin.com/in/johndoe"
                        )
                    
                    # Education Section with dynamic fields
                    with gr.Group(elem_classes=["form-section"]):
                        gr.HTML("<h3>🎓 Education</h3>")
                        
                        # Education entries (supporting multiple degrees)
                        education_institution = gr.Textbox(
                            label="Institution/University *", 
                            placeholder="University of Technology, Harvard University",
                            lines=2
                        )
                        
                        with gr.Row():
                            education_degree = gr.Textbox(
                                label="Degree/Program", 
                                placeholder="Bachelor of Science in Computer Science, Master of Business Administration",
                                lines=2
                            )
                            education_start = gr.Textbox(
                                label="Start Year", 
                                placeholder="2018, 2022",
                                lines=2
                            )
                            education_end = gr.Textbox(
                                label="End Year", 
                                placeholder="2022, 2024",
                                lines=2
                            )
                        
                        with gr.Row():
                            education_gpa = gr.Textbox(
                                label="GPA/CGPA", 
                                placeholder="3.8/4.0, 3.9/4.0"
                            )
                            education_percentage = gr.Textbox(
                                label="Percentage (10th/12th)", 
                                placeholder="85, 92"
                            )
                        
                        education_coursework = gr.Textbox(
                            label="Relevant Coursework", 
                            placeholder="Data Structures, Algorithms, Database Management, Machine Learning",
                            lines=2
                        )
                    
                    # Skills Section
                    with gr.Group(elem_classes=["form-section"]):
                        gr.HTML("<h3>💼 Skills & Technologies</h3>")
                        user_skills = gr.Textbox(
                            label="Technical Skills *", 
                            placeholder="Python, JavaScript, React, Node.js, SQL, AWS, Docker, Git, Agile, Project Management",
                            lines=4,
                            info="Separate skills with commas. Include both technical and soft skills."
                        )
                    
                    # Experience Section
                    with gr.Group(elem_classes=["form-section"]):
                        gr.HTML("<h3>💼 Professional Experience</h3>")
                        
                        experience_company = gr.Textbox(
                            label="Company Names", 
                            placeholder="Tech Solutions Inc., Innovation Labs",
                            lines=2
                        )
                        
                        with gr.Row():
                            experience_position = gr.Textbox(
                                label="Job Titles", 
                                placeholder="Software Developer, Senior Frontend Developer",
                                lines=2
                            )
                            experience_location = gr.Textbox(
                                label="Work Locations", 
                                placeholder="San Francisco, CA; Remote",
                                lines=2
                            )
                        
                        with gr.Row():
                            experience_start_date = gr.Textbox(
                                label="Start Dates", 
                                placeholder="June 2022, January 2024"
                            )
                            experience_end_date = gr.Textbox(
                                label="End Dates", 
                                placeholder="December 2023, Present"
                            )
                        
                        experience_responsibilities = gr.Textbox(
                            label="Job Responsibilities & Achievements", 
                            placeholder="""• Developed full-stack web applications using React and Node.js, serving 10,000+ users
• Implemented CI/CD pipelines reducing deployment time by 40%
• Led a team of 3 developers on critical project delivery
• Optimized database queries improving application performance by 60%""",
                            lines=6,
                            info="Use bullet points. Focus on achievements with quantifiable results."
                        )
                    
                    # Projects Section
                    with gr.Group(elem_classes=["form-section"]):
                        gr.HTML("<h3>🚀 Projects</h3>")
                        
                        project_name = gr.Textbox(
                            label="Project Names", 
                            placeholder="E-commerce Platform, Task Management App, Portfolio Website",
                            lines=2
                        )
                        
                        with gr.Row():
                            project_date = gr.Textbox(
                                label="Project Dates", 
                                placeholder="2023, 2024, 2024"
                            )
                            project_technologies = gr.Textbox(
                                label="Technologies Used", 
                                placeholder="React, Node.js, MongoDB; Python, Flask, PostgreSQL; HTML, CSS, JavaScript",
                                lines=2
                            )
                        
                        project_description = gr.Textbox(
                            label="Project Descriptions", 
                            placeholder="""Built a full-stack e-commerce platform with payment integration and admin dashboard
Developed a collaborative task management application with real-time updates
Created a responsive portfolio website showcasing development projects""",
                            lines=4
                        )
                        
                        project_url = gr.Textbox(
                            label="Project URLs/GitHub Links", 
                            placeholder="https://github.com/johndoe/ecommerce, https://taskapp.com, https://johndoe.dev"
                        )
                    
                    # Additional Information
                    with gr.Group(elem_classes=["form-section"]):
                        gr.HTML("<h3>📋 Additional Information</h3>")
                        
                        with gr.Row():
                            certifications = gr.Textbox(
                                label="Certifications", 
                                placeholder="AWS Certified Developer, Google Cloud Professional, Scrum Master Certified",
                                lines=2
                            )
                            languages = gr.Textbox(
                                label="Languages", 
                                placeholder="English (Native), Spanish (Fluent), French (Intermediate)"
                            )
                        
                        interests = gr.Textbox(
                            label="Interests/Hobbies", 
                            placeholder="Open Source Contributing, Photography, Marathon Running, Chess"
                        )
                
                with gr.Column(scale=1):
                    # Template Selection
                    with gr.Group(elem_classes=["form-section"]):
                        gr.HTML("<h3>🎨 Choose Resume Template</h3>")
                        
                        template_choice = gr.Radio(
                            choices=[
                                ("Modern Professional", "modern"),
                                ("Classic Executive", "classic"), 
                                ("Creative Portfolio", "creative"),
                                ("Minimalist", "minimal"),
                                ("Tech Professional", "tech")
                            ],
                            value="modern",
                            label="Select Template Style"
                        )
                        
                        gr.HTML("""
                        <div class="info-card">
                            <h4>Template Features:</h4>
                            <ul>
                                <li><strong>Modern:</strong> Clean design with accent colors</li>
                                <li><strong>Classic:</strong> Traditional, executive-style layout</li>
                                <li><strong>Creative:</strong> Vibrant design for creative roles</li>
                                <li><strong>Minimal:</strong> Simple, content-focused layout</li>
                                <li><strong>Tech:</strong> Optimized for technical positions</li>
                            </ul>
                        </div>
                        """)
                    
                    # Target Position Section
                    with gr.Group(elem_classes=["form-section"]):
                        gr.HTML("<h3>🎯 Target Position</h3>")
                        
                        desired_position = gr.Textbox(
                            label="Desired Job Title *", 
                            placeholder="Senior Software Engineer, Product Manager, Data Scientist"
                        )
                        
                        job_description = gr.Textbox(
                            label="Job Description (for ATS optimization)", 
                            placeholder="Paste the complete job description here. This helps optimize your resume for ATS systems and ensures better keyword matching.",
                            lines=8,
                            info="Adding job description significantly improves ATS score and keyword matching."
                        )
                    
                    # Generate Button
                    generate_button = gr.Button(
                        "🚀 Generate Optimized Resume", 
                        elem_classes=["primary"],
                        variant="primary",
                        size="lg"
                    )
                    
                    # AI Enhancement Toggle
                    with gr.Group(elem_classes=["form-section"]):
                        gr.HTML("<h3>🤖 AI Enhancement Options</h3>")
                        
                        enhance_content = gr.Checkbox(
                            label="Auto-enhance content with AI",
                            value=True,
                            info="Uses AI to improve descriptions and add relevant keywords"
                        )
                        
                        include_summary = gr.Checkbox(
                            label="Generate professional summary",
                            value=True,
                            info="Creates a tailored professional summary based on job description"
                        )
                        
                        keyword_optimization = gr.Checkbox(
                            label="Optimize for ATS keywords",
                            value=True,
                            info="Automatically includes relevant keywords from job description"
                        )
            
            # Results Section
            with gr.Group(elem_classes=["form-section"], visible=False) as results_container:
                gr.HTML("<h3>📊 Resume Generated Successfully!</h3>")
                
                with gr.Row():
                    with gr.Column(scale=1):
                        score_display = gr.HTML(elem_classes=["score-display"])
                        
                        # Download and action buttons
                        with gr.Row():
                            resume_download = gr.File(
                                label="📄 Download Your Resume",
                                file_count="single"
                            )
                            
                        with gr.Row():
                            regenerate_btn = gr.Button(
                                "🔄 Regenerate with Different Template",
                                elem_classes=["secondary"],
                                variant="secondary"
                            )
                            
                            improve_btn = gr.Button(
                                "✨ Get Improvement Suggestions",
                                elem_classes=["primary"],
                                variant="primary"
                            )
                    
                    with gr.Column(scale=2):
                        # Stats cards
                        with gr.Row():
                            gr.HTML("""
                            <div class="stats-card">
                                <div class="stats-number" id="keyword-match">0%</div>
                                <div class="stats-label">Keyword Match</div>
                            </div>
                            """)
                            
                            gr.HTML("""
                            <div class="stats-card">
                                <div class="stats-number" id="sections-count">0</div>
                                <div class="stats-label">Resume Sections</div>
                            </div>
                            """)
                            
                            gr.HTML("""
                            <div class="stats-card">
                                <div class="stats-number" id="total-words">0</div>
                                <div class="stats-label">Total Words</div>
                            </div>
                            """)
                        
                        # AI Suggestions Section
                        with gr.Group(elem_classes=["suggestion-box"]):
                            gr.HTML("<h4>💡 AI-Powered Suggestions</h4>")
                            ai_suggestions_output = gr.Markdown(
                                value="Your AI suggestions will appear here after generating the resume.",
                                elem_classes=["suggestion-content"]
                            )
        
        # Tab 2: Enhanced ATS Analyzer
        with gr.Tab("🔍 ATS Score Analyzer", elem_id="ats-tab"):
            with gr.Column(elem_classes=["container"]):
                gr.HTML("""
                <div class="info-card">
                    <h3>🎯 ATS Compatibility Analyzer</h3>
                    <p>Upload your existing resume and analyze how well it matches with job descriptions. 
                    Get detailed feedback on keyword optimization, formatting, and content structure.</p>
                </div>
                """)
                
                with gr.Row():
                    with gr.Column(scale=1):
                        gr.HTML("""
                        <div class="file-upload">
                            <input type="file" accept=".pdf,.docx" />
                            <label class="file-upload-label">
                                <div class="file-upload-icon">📄</div>
                                <strong>Upload Your Resume</strong>
                                <span>Drag & drop or click to browse</span>
                                <small>Supports PDF and DOCX files</small>
                            </label>
                        </div>
                        """)
                        
                        uploaded_resume = gr.File(
                            label="Upload Resume", 
                            file_types=[".docx", ".pdf"],
                            elem_classes=["file-upload-input"]
                        )
                        
                        # Quick analysis options
                        with gr.Group(elem_classes=["form-section"]):
                            gr.HTML("<h4>⚙️ Analysis Options</h4>")
                            
                            detailed_analysis = gr.Checkbox(
                                label="Detailed keyword analysis",
                                value=True
                            )
                            
                            competitor_analysis = gr.Checkbox(
                                label="Compare with industry standards",
                                value=False
                            )
                            
                            format_check = gr.Checkbox(
                                label="Check formatting and structure",
                                value=True
                            )
                    
                    with gr.Column(scale=2):
                        job_description_input = gr.Textbox(
                            label="Target Job Description", 
                            placeholder="Paste the complete job description here for accurate ATS analysis...",
                            lines=12,
                            info="The more detailed the job description, the more accurate the ATS analysis will be."
                        )
                        
                        # Industry selector
                        industry_selector = gr.Dropdown(
                            choices=[
                                "Technology/Software",
                                "Finance/Banking",
                                "Healthcare",
                                "Marketing/Advertising",
                                "Sales",
                                "Education",
                                "Manufacturing",
                                "Consulting",
                                "Other"
                            ],
                            label="Industry/Field",
                            value="Technology/Software",
                            info="Helps provide industry-specific optimization suggestions"
                        )
                
                analyze_button = gr.Button(
                    "🔍 Analyze ATS Compatibility", 
                    elem_classes=["primary"],
                    variant="primary",
                    size="lg"
                )
                
                # Analysis Results Section
                with gr.Group(elem_classes=["form-section"], visible=False) as analysis_results:
                    gr.HTML("<h3>📊 ATS Analysis Results</h3>")
                    
                    with gr.Row():
                        with gr.Column(scale=1):
                            ats_score_display = gr.HTML(elem_classes=["score-display"])
                            
                            # Score breakdown
                            gr.HTML("""
                            <div class="info-card">
                                <h4>Score Breakdown</h4>
                                <div class="progress-bar">
                                    <div class="progress-fill" style="width: 0%" id="keyword-progress"></div>
                                </div>
                                <p>Keyword Match: <span id="keyword-percentage">0%</span></p>
                                
                                <div class="progress-bar">
                                    <div class="progress-fill" style="width: 0%" id="format-progress"></div>
                                </div>
                                <p>Format Score: <span id="format-percentage">0%</span></p>
                                
                                <div class="progress-bar">
                                    <div class="progress-fill" style="width: 0%" id="content-progress"></div>
                                </div>
                                <p>Content Quality: <span id="content-percentage">0%</span></p>
                            </div>
                            """)
                        
                        with gr.Column(scale=2):
                            # Detailed Analysis Tabs
                            with gr.Tabs():
                                with gr.Tab("💡 Recommendations"):
                                    ai_suggestion_result = gr.Markdown(
                                        elem_classes=["suggestion-content"]
                                    )
                                
                                with gr.Tab("🔍 Keyword Analysis"):
                                    keyword_analysis_result = gr.HTML()
                                
                                with gr.Tab("📋 Missing Elements"):
                                    missing_elements_result = gr.HTML()
                                
                                with gr.Tab("✅ Strengths"):
                                    strengths_result = gr.HTML()
        
        # Tab 3: Enhanced AI Career Coach
        with gr.Tab("💬 AI Career Coach", elem_id="chat-tab"):
            with gr.Column(elem_classes=["container"]):
                gr.HTML("""
                <div class="info-card">
                    <h3>🤖 Your Personal AI Career Coach</h3>
                    <p>Get personalized career advice, interview preparation, resume tips, and job search strategies. 
                    Our AI coach is trained on the latest industry trends and hiring practices.</p>
                </div>
                """)
                
                # Quick action buttons
                with gr.Row():
                    gr.HTML("""
                    <div class="info-card">
                        <h4>💡 Popular Questions</h4>
                        <div style="display: flex; flex-wrap: wrap; gap: 8px; margin-top: 12px;">
                            <button class="secondary" onclick="document.getElementById('chat-input').value='How do I explain a career gap in my resume?'">Career Gaps</button>
                            <button class="secondary" onclick="document.getElementById('chat-input').value='What are the best interview questions to ask?'">Interview Prep</button>
                            <button class="secondary" onclick="document.getElementById('chat-input').value='How do I negotiate salary effectively?'">Salary Negotiation</button>
                            <button class="secondary" onclick="document.getElementById('chat-input').value='What skills are in demand for my industry?'">Skill Development</button>
                        </div>
                    </div>
                    """)
                
                # Chat interface
                chat_history = gr.Chatbot(
                    elem_id="chat-container",
                    elem_classes=["chatbot-container"],
                    height=500,
                    show_label=False,
                    avatar_images=["👤", "🤖"],
                    type="messages"
                )
                
                with gr.Row():
                    chat_input = gr.Textbox(
                        label="Ask your career question", 
                        placeholder="Type your career question here... (e.g., 'How do I transition to a tech career?')",
                        lines=2,
                        elem_id="chat-input",
                        show_label=False
                    )
                    
                    with gr.Column(scale=0):
                        chat_button = gr.Button(
                            "Send", 
                            elem_classes=["primary"],
                            variant="primary"
                        )
                        
                        clear_chat = gr.Button(
                            "Clear", 
                            elem_classes=["secondary"],
                            variant="secondary"
                        )
                
                # Chat suggestions
                with gr.Row():
                    gr.HTML("""
                    <div class="suggestion-box">
                        <h4>💬 Conversation Starters</h4>
                        <ul>
                            <li>"I'm switching careers from [current field] to [target field]. What should I focus on?"</li>
                            <li>"How do I make my resume stand out for [specific role]?"</li>
                            <li>"What are the most important skills for [industry] in 2024?"</li>
                            <li>"How do I prepare for behavioral interview questions?"</li>
                            <li>"What's the best way to follow up after an interview?"</li>
                        </ul>
                    </div>
                    """)
        
        # Tab 4: Resume Templates & Examples
        with gr.Tab("📚 Templates & Examples", elem_id="templates-tab"):
            with gr.Column(elem_classes=["container"]):
                gr.HTML("""
                <div class="info-card">
                    <h3>📋 Resume Templates & Industry Examples</h3>
                    <p>Explore different resume formats and see examples from various industries. 
                    Get inspiration and understand what works best for your field.</p>
                </div>
                """)
                
                # Template gallery
                with gr.Row():
                    for template_key, template_info in RESUME_TEMPLATES.items():
                        with gr.Column():
                            gr.HTML(f"""
                            <div class="template-card" onclick="selectTemplate('{template_key}')">
                                <h4>{template_info['name']}</h4>
                                <p>{template_info['description']}</p>
                                <div style="margin-top: 12px;">
                                    <button class="secondary">Preview Template</button>
                                </div>
                            </div>
                            """)
                
                # Industry-specific examples
                with gr.Group(elem_classes=["form-section"]):
                    gr.HTML("<h3>🏭 Industry-Specific Examples</h3>")
                    
                    industry_examples = gr.Dropdown(
                        choices=[
                            "Software Engineering",
                            "Data Science",
                            "Product Management", 
                            "Marketing",
                            "Finance",
                            "Healthcare",
                            "Education",
                            "Sales"
                        ],
                        label="Select Industry for Examples",
                        value="Software Engineering"
                    )
                    
                    example_display = gr.HTML("""
                    <div class="info-card">
                        <h4>Software Engineering Resume Tips</h4>
                        <ul>
                            <li><strong>Technical Skills:</strong> List programming languages, frameworks, and tools prominently</li>
                            <li><strong>Projects:</strong> Include 2-3 significant projects with GitHub links</li>
                            <li><strong>Quantify Impact:</strong> Use metrics like "improved performance by 40%" or "reduced load time by 2 seconds"</li>
                            <li><strong>Keywords:</strong> Include relevant technologies mentioned in job descriptions</li>
                            <li><strong>Format:</strong> Keep it clean and scannable, avoid fancy graphics</li>
                        </ul>
                    </div>
                    """)
    
    # JavaScript for enhanced interactivity
    gr.HTML("""
    <script>
    function selectTemplate(templateKey) {
        // Update template selection
        console.log('Selected template:', templateKey);
        // You can add more interactive features here
    }
    
    // Add smooth scrolling to form sections
    document.addEventListener('DOMContentLoaded', function() {
        const formSections = document.querySelectorAll('.form-section');
        formSections.forEach(section => {
            section.addEventListener('click', function() {
                this.scrollIntoView({ behavior: 'smooth', block: 'nearest' });
            });
        });
    });
    
    // Auto-save form data to localStorage
    function autoSaveForm() {
        const inputs = document.querySelectorAll('input, textarea, select');
        inputs.forEach(input => {
            input.addEventListener('change', function() {
                localStorage.setItem(this.id || this.name, this.value);
            });
        });
    }
    
    // Load saved form data
    function loadSavedData() {
        const inputs = document.querySelectorAll('input, textarea, select');
        inputs.forEach(input => {
            const savedValue = localStorage.getItem(input.id || input.name);
            if (savedValue) {
                input.value = savedValue;
            }
        });
    }
    
    // Initialize auto-save
    setTimeout(autoSaveForm, 1000);
    setTimeout(loadSavedData, 500);
    </script>
    """)
    
    # Enhanced event handlers
    def process_and_generate_resume(full_name, user_email, contact_number, location, linkedin,
                                   education_degree, education_institution, education_start, 
                                   education_end, education_gpa, education_percentage, education_coursework,
                                   user_skills, experience_company, experience_position, experience_start_date,
                                   experience_end_date, experience_location, experience_responsibilities,
                                   project_name, project_date, project_description, project_technologies,
                                   project_url, certifications, languages, interests, desired_position,
                                   job_description, template_choice):
        
        # Process the form data
        resume_data = process_resume_form(
            full_name, user_email, contact_number, location, linkedin,
            [education_degree], [education_institution], [education_start], [education_end],
            [education_gpa], [education_percentage], [education_coursework],
            user_skills, [experience_company], [experience_position], [experience_start_date],
            [experience_end_date], [experience_location], [experience_responsibilities],
            [project_name], [project_date], [project_description], [project_technologies],
            [project_url], certifications, languages, interests, desired_position, job_description
        )
        
        # Generate the resume
        filename, ats_text, ai_suggestions, score = generate_resume(resume_data, job_description, template_choice)
        
        # Update score display
        score_class = update_score_class(score)
        score_html = f"""
        <div class="{score_class}">
            <div style="position: relative; z-index: 1;">
                {score}%
                <div style="font-size: 1rem; margin-top: 8px; opacity: 0.9;">ATS Compatibility</div>
            </div>
        </div>
        """
        
        return (
            gr.update(visible=True),  # Show results container
            score_html,               # Score display
            filename,                 # Resume file
            ai_suggestions           # AI suggestions
        )
    
    # Connect the generate button to the enhanced function
    generate_button.click(
        process_and_generate_resume,
        inputs=[
            full_name, user_email, contact_number, location, linkedin,
            education_degree, education_institution, education_start, education_end,
            education_gpa, education_percentage, education_coursework,
            user_skills, experience_company, experience_position, experience_start_date,
            experience_end_date, experience_location, experience_responsibilities,
            project_name, project_date, project_description, project_technologies,
            project_url, certifications, languages, interests, desired_position,
            job_description, template_choice
        ],
        outputs=[results_container, score_display, resume_download, ai_suggestions_output]
    )
    
    # Enhanced ATS analyzer function
    def enhanced_analyze_resume(uploaded_resume, job_description, industry, detailed_analysis, format_check):
        if not uploaded_resume or not job_description.strip():
            return "⚠️ Please upload a resume and provide a job description.", "", 0.0
        
        # Existing analysis logic
        ats_text, suggestions, score = analyze_uploaded_resume(uploaded_resume, job_description)
        
        # Enhanced score display
        score_class = update_score_class(score)
        score_html = f"""
        <div class="{score_class}">
            <div style="position: relative; z-index: 1;">
                {score}%
                <div style="font-size: 1rem; margin-top: 8px; opacity: 0.9;">ATS Score</div>
            </div>
        </div>
        """
        
        return (
            gr.update(visible=True),  # Show analysis results
            score_html,               # Score display
            suggestions              # AI suggestions
        )
    
    # Connect ATS analyzer
    analyze_button.click(
        enhanced_analyze_resume,
        inputs=[uploaded_resume, job_description_input, industry_selector, detailed_analysis, format_check],
        outputs=[analysis_results, ats_score_display, ai_suggestion_result]
    )
    
    # Enhanced chat functionality
    def enhanced_chatbot(user_query, history):
        if not user_query.strip():
            return history, ""
        
        # Enhanced prompt for career coaching
        enhanced_prompt = f"""
        You are an expert career coach and resume consultant with 15+ years of experience helping professionals advance their careers. 
        
        User Question: {user_query}
        
        Please provide comprehensive, actionable advice that is:
        - Specific and practical
        - Based on current industry trends
        - Tailored to the user's apparent career level and goals
        - Supportive and encouraging
        
        If the question is about resumes, include specific formatting and content suggestions.
        If about interviews, provide example questions and answers.
        If about career transitions, outline concrete steps and timelines.
        """
        
        try:
            model = genai.GenerativeModel("gemini-1.5-flash")
            response = model.generate_content(enhanced_prompt)
            
            # Format for new message type
            new_message = {"role": "assistant", "content": response.text}
            history.append({"role": "user", "content": user_query})
            history.append(new_message)
            
        except Exception as e:
            error_message = {"role": "assistant", "content": f"I apologize, but I encountered an error: {str(e)}. Please try again."}
            history.append({"role": "user", "content": user_query})
            history.append(error_message)
        
        return history, ""
    
    # Connect chat functionality
    chat_input.submit(enhanced_chatbot, inputs=[chat_input, chat_history], outputs=[chat_history, chat_input])
    chat_button.click(enhanced_chatbot, inputs=[chat_input, chat_history], outputs=[chat_history, chat_input])
    
    # Clear chat function
    def clear_chat_history():
        return []
    
    clear_chat.click(clear_chat_history, outputs=[chat_history])

# Launch the enhanced application
if __name__ == "__main__":
    iface.launch(
        share=False,  # Disable share link to avoid connection issues
        server_name="127.0.0.1",  # Use localhost instead of 0.0.0.0
        server_port=7860,
        show_api=False,
        favicon_path=None,
        ssl_verify=False,
        quiet=False,
        inbrowser=True  # Automatically open browser
    )
