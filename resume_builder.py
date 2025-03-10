import os
import spacy
import fitz  # PyMuPDF for PDF text extraction
import gradio as gr
import google.generativeai as genai
from docx import Document

# Configure Google Gemini API
genai.configure(api_key="AIzaSyBIp7DNNfyAFxUQKiycBdRadMfvDZ06wmU")  # Replace with your API key

# Load NLP model for ATS Score Calculation
nlp = spacy.load("en_core_web_sm")

# ✅ Function to Interact with Google Gemini API
def chatbot_interface(user_query, history=[]):
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(user_query)
    history.append((user_query, response.text))
    return history, ""  # Clear input after sending message

# ✅ Function to Calculate ATS Score
def calculate_ats_score(resume_text, job_description):
    if not job_description.strip():
        return "⚠️ Please provide a job description to calculate ATS score.", 0.0
    
    resume_doc = nlp(resume_text.lower())
    job_doc = nlp(job_description.lower())
    resume_words = {token.lemma_ for token in resume_doc if token.is_alpha}
    job_words = {token.lemma_ for token in job_doc if token.is_alpha}
    
    if not job_words:
        return "⚠️ Job description is too short to analyze.", 0.0
    
    match_score = (len(resume_words & job_words) / len(job_words)) * 100
    return f"📊 ATS Score: {round(match_score, 2)}%", round(match_score, 2)

# ✅ AI Suggestions for Resume Improvement
def ai_suggestions(resume_text, job_description):
    if not job_description.strip():
        return "⚠️ Please provide a job description for AI suggestions."
    
    prompt = f"""
    Analyze this resume and provide suggestions for improvement based on the given job description.

    Resume: {resume_text}
    
    Job Description: {job_description}
    
    Provide a detailed AI-generated suggestion.
    """
    
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(prompt)
    return response.text

# ✅ Function to Generate Resume
def resume_chatbot(full_name, user_email, contact_number, user_education, user_skills, user_experience, job_position, job_description):
    if not full_name or not user_email or not contact_number or not user_education or not user_skills or not user_experience or not job_position:
        return None, "⚠️ Please fill in all fields to generate a resume.", ""
    
    doc = Document()
    doc.add_heading('Professional Resume', level=1)
    doc.add_paragraph(f"Name: {full_name}")
    doc.add_paragraph(f"Email: {user_email}")
    doc.add_paragraph(f"Phone: {contact_number}")
    doc.add_paragraph(f"Education: {user_education}")
    doc.add_paragraph(f"Skills: {user_skills}")
    doc.add_paragraph(f"Experience: {user_experience}")
    doc.add_paragraph(f"Applying for Role: {job_position}")

    ats_text, _ = calculate_ats_score(user_skills, job_description)
    doc.add_paragraph(ats_text)
    
    resume_file = f"resume_{full_name.replace(' ', '_')}.docx"
    doc.save(resume_file)

    ai_suggestion_text = ai_suggestions(f"Skills: {user_skills}, Experience: {user_experience}", job_description)
    
    return resume_file, ats_text, ai_suggestion_text

# ✅ Function to Analyze Uploaded Resume
def analyze_uploaded_resume(uploaded_resume, job_description):
    if not uploaded_resume or not job_description.strip():
        return "⚠️ Please upload a resume and provide a job description.", None
    
    resume_text = ""
    file_ext = uploaded_resume.name.split(".")[-1].lower()
    
    try:
        if file_ext == "docx":
            doc = Document(uploaded_resume.name)
            resume_text = "\n".join([para.text for para in doc.paragraphs])
        elif file_ext == "pdf":
            with fitz.open(uploaded_resume.name) as pdf_doc:
                resume_text = "\n".join([page.get_text("text") for page in pdf_doc])
        else:
            return "❌ Unsupported file format! Please upload a DOCX or PDF file.", None
        
        ats_text, _ = calculate_ats_score(resume_text, job_description)
        ai_suggestion_text = ai_suggestions(resume_text, job_description)
        return ats_text, ai_suggestion_text
    except Exception as e:
        return f"❌ Error processing resume: {str(e)}", None

# ✅ Gradio UI
with gr.Blocks(css="""
    body { background-color: #121212; color: white; font-family: Arial; }
    .gradio-container { width: 90%; max-width: 800px; margin: auto; }
    .gradio-button { background-color: #FF9800; color: black; border-radius: 8px; transition: transform 0.1s ease-in-out; }
    .gradio-button:active { transform: scale(0.95); }
    .gradio-textbox, .gradio-file { background-color: #1E1E1E; color: white; }
    @keyframes fadeIn { from { opacity: 0; } to { opacity: 1; } }
    .tab-content { animation: fadeIn 0.5s ease-in-out; }
""") as iface:
    
    gr.Markdown("""<h1 style='text-align: center; color: #FFA500;'>📝 AI Resume Builder 🚀</h1>""")

    with gr.Tab("📄 Resume Generator", elem_classes=["tab-content"]):
        full_name, user_email, contact_number = gr.Textbox(label="Full Name"), gr.Textbox(label="Email"), gr.Textbox(label="Phone Number")
        user_education, user_skills, user_experience = gr.Textbox(label="Education"), gr.Textbox(label="Skills"), gr.Textbox(label="Experience")
        job_position, job_description = gr.Textbox(label="Job Role"), gr.Textbox(label="Job Description")
        generate_button = gr.Button("🚀 Generate Resume")
        resume_download, ats_score_output, ai_suggestions_output = gr.File(label="📄 Download DOCX"), gr.Text(label="📊 ATS Score"), gr.Text(label="💡 AI Suggestions")
        
        generate_button.click(
            resume_chatbot,   
            inputs=[full_name, user_email, contact_number, user_education, user_skills, user_experience, job_position, job_description],
            outputs=[resume_download, ats_score_output, ai_suggestions_output]
        )
    
    with gr.Tab("📊 ATS Score Analyzer", elem_classes=["tab-content"]):
        uploaded_resume, job_description_input = gr.File(label="📄 Upload Resume"), gr.Textbox(label="Job Description")
        analyze_button = gr.Button("🔍 Analyze ATS Score")
        ats_result, ai_suggestion_result =  gr.Text(label="📊 ATS Score"), gr.Text(label="💡 AI Suggestions")
        analyze_button.click(analyze_uploaded_resume, [uploaded_resume, job_description_input], [ats_result, ai_suggestion_result])

    with gr.Tab("💬 AI Chat Assist", elem_classes=["tab-content"]):
        chat_history = gr.Chatbot(label="Chat", elem_id="chat-container")
        chat_input = gr.Textbox(label="Ask AI Assistant", interactive=True)
        chat_button = gr.Button("🤖 Chat")
        chat_input.submit(chatbot_interface, inputs=[chat_input, chat_history], outputs=[chat_history, chat_input])
        chat_button.click(chatbot_interface, inputs=[chat_input, chat_history], outputs=[chat_history, chat_input])

iface.launch(share=True)
